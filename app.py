import customtkinter
from PIL import Image
from tkinter import filedialog, messagebox
import os
import time  
from PIL import ImageGrab  

import io  
import threading

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

from image_caption import caption_image, caption_image_from_bytes
from docx.enum.style import WD_STYLE_TYPE
from docx import Document  
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
from settings import load_settings, save_settings, DEFAULT_SETTINGS
import sys
import hashlib
# add near top of file (once)

import platform

from datetime import datetime


# --- NEW: Imports to fix the .exe path ---
SETTINGS = load_settings()
COM_LOCK = threading.Lock()
USE_WORD_AUTOMATION = False

# --- Caption cooldown ---
LAST_CAPTION_TIME = 0
CAPTION_COOLDOWN_SECONDS = 5.0  # Wait 5 seconds between captions
LAST_IMAGE_HASH = None

if platform.system() == "Windows":
    try:
        import win32com.client
        USE_WORD_AUTOMATION = True
        print("✔ Word automation available (pywin32).")
    except ImportError:
        print("⚠ pywin32 not installed — falling back to python-docx captions.")
else:
    print("⚠ Non-Windows OS — using python-docx captions.")


# --- THEME 3 COLORS (User's custom theme) ---
COLOR_WINDOW_BG = "#E3E4F0"
COLOR_BUTTON = "#A9B5DF"
COLOR_BUTTON_HOVER = "#7886C7"
COLOR_TEXT = "#2D336B"
COLOR_TEXT_IDLE = "#60688A"
COLOR_TEXT_ACTIVE = "#2D336B"
COLOR_TEXT_ERROR = "#D9534F"
COLOR_TEXT_DISABLED = "#60688A"

# --- Global Variables ---
selected_doc_path = None
monitoring = False

selected_folder_path = None
monitoring_mode = "Clipboard" 
folder_observer = None 

# --- NEW: Function to find assets (icons) in the .exe ---
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        # _MEIPASS not an attribute, so we're not in a PyInstaller bundle
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# --- Thread-safe error handlers ---
def handle_permission_error():
    messagebox.showerror(
        "File Error",
        f"Could not save to {os.path.basename(selected_doc_path)}.\n\n"
        "Please make sure the file is CLOSED in Microsoft Word."
    )
    stop_monitoring() 

def handle_generic_error(error_message):
    messagebox.showerror("Error", f"An unknown error occurred: {error_message}")
    print(f"Error: {error_message}")
    stop_monitoring() 

def add_caption_python_docx(doc, caption_text, prefix="Figure", numbering="decimal"):
    """
    Creates a numbered caption using SEQ fields and applies Caption style.
    numbering: "decimal", "loweralpha", "lowerroman"
    """
    # Ensure style
    try:
        doc.styles["Caption"]
    except KeyError:
        style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(SETTINGS.get("caption_style", {}).get("size_pt", 9))
        style.font.italic = SETTINGS.get("caption_style", {}).get("italic", True)

    paragraph = doc.add_paragraph(style="Caption")

    # Build SEQ instr with prefix and numbering mapping
    # Example: SEQ Figure \* ARABIC
    numbering_map = {
        "decimal": r"\* Arabic",
        "loweralpha": r"\* ALPHABETIC",
        "lowerroman": r"\* roman"
    }
    num_token = numbering_map.get(numbering, r"\* Arabic")

    instr = f'SEQ {prefix} {num_token}'
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)

    run = paragraph.add_run()
    run._r.append(fld)

    paragraph.add_run(f": {caption_text}")
    return paragraph



# Helper: choose insertion method (pywin vs python-docx)
def insert_real_caption(doc_path, caption_text, prefix="Figure"):
    """
    Insert a REAL Word caption using COM automation.
    Falls back to styled paragraph only if COM fails.
    Logs which method was used.
    """
    
    if not USE_WORD_AUTOMATION:
        print("⚠ Word automation not available. Using styled paragraph instead.")
        doc = Document(doc_path)
        add_caption_python_docx(doc, caption_text, prefix=prefix, numbering=SETTINGS.get("numbering_format", "decimal"))
        doc.save(doc_path)
        print("📝 Caption inserted as: STYLED PARAGRAPH (fallback)")
        return

    # Try COM automation first
    def _try_com_caption():
        word = None
        doc = None
        try:
            with COM_LOCK:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                # Open document
                doc = word.Documents.Open(doc_path)
                
                # Get total number of InlineShapes (images)
                if doc.InlineShapes.Count == 0:
                    print("⚠ No images found in document")
                    return False
                
                # Select the last image
                last_image = doc.InlineShapes(doc.InlineShapes.Count)
                last_image.Range.Select()
                
                # Now insert caption - Word's InsertCaption works on selected image
                selection = word.Selection
                selection.InsertCaption(
                    Label=prefix,
                    Title=f": {caption_text}",  # Add colon and space here
                    Position=1,
                    ExcludeLabel=False
                )
                
                print("✅ Caption inserted as: REAL WORD CAPTION (COM automation)")
                doc.Save()
                return True
                
        except Exception as e:
            print(f"⚠ COM caption failed: {e}")
            return False
        finally:
            try:
                if doc:
                    doc.Close(False)
                if word:
                    word.Quit()
            except:
                pass

    # Try COM first
    success = _try_com_caption()
    
    # Fallback to styled paragraph if COM fails
    if not success:
        try:
            doc = Document(doc_path)
            add_caption_python_docx(doc, caption_text, prefix=prefix, numbering=SETTINGS.get("numbering_format", "decimal"))
            doc.save(doc_path)
            print("📝 Caption inserted as: STYLED PARAGRAPH (COM failed, using fallback)")
        except Exception as e:
            print(f"❌ Caption insertion failed completely: {e}")


def insert_image_at_codeword(doc, image_stream):
    codeword = SETTINGS.get("insertion_codeword", "").strip()
    if not codeword:
        return None

    for para in doc.paragraphs:
        if codeword in para.text:
            para.text = para.text.replace(codeword, "", 1)

            run = para.add_run()
            run.add_picture(image_stream, width=Inches(6.0))

            return para   # 🔥 RETURN THE IMAGE PARAGRAPH

    return None

# --- NEW: Image Compression Helper ---
def compress_image(image_bytes):
    """
    Compresses image bytes (JPEG) and optionally resizes.
    Returns (compressed_bytes, is_compressed_flag)
    """
    if not SETTINGS.get("compression_enabled", False):
        return image_bytes, False

    try:
        img = Image.open(io.BytesIO(image_bytes))
        
        # Convert to RGB for JPEG
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")

        out_io = io.BytesIO()
        quality = int(SETTINGS.get("compression_quality", 85))
        img.save(out_io, format="JPEG", quality=quality, optimize=True)
        return out_io.getvalue(), True
    except Exception as e:
        print(f"⚠️ Compression failed: {e}")
        return image_bytes, False

# --- MODIFIED: Core Logic (replacement for your existing function) ---
def add_image_to_doc(image_input, input_type="data"):
    global monitoring, selected_doc_path, LAST_CAPTION_TIME
    
    if not monitoring:
        return

    print(f"Adding image to {selected_doc_path}...")

    try:
        # Get image bytes
        if input_type == "data":
            image_bytes = image_input
        else:
            with open(image_input, "rb") as f:
                image_bytes = f.read()

        image_stream = io.BytesIO(image_bytes)

        # --- COMPRESSION STEP ---
        is_compressed = False
        compressed_bytes, is_compressed = compress_image(image_bytes)
        
        # Use compressed bytes for doc insertion AND saving copy (if successful)
        final_image_bytes = compressed_bytes
        final_image_stream = io.BytesIO(final_image_bytes)
        
        # Determine extension
        ext = ".jpg" if is_compressed else ".png"

        # --- Save Copy to Folder (New Feature) ---
        if SETTINGS.get("save_copy_enabled", False):
            folder = SETTINGS.get("save_copy_folder", "")
            if folder and os.path.exists(folder):
                try:
                    # 1. Create Date-based subfolder (YYYY-MM-DD)
                    now_ts = datetime.now()
                    date_folder_name = now_ts.strftime("%Y-%m-%d")
                    target_dir = os.path.join(folder, date_folder_name)
                    os.makedirs(target_dir, exist_ok=True)

                    # 2. Find next available incremental filename (screenshot_001.png/jpg)
                    counter = 1
                    while True:
                        filename = f"screenshot_{counter:03d}{ext}"
                        save_path = os.path.join(target_dir, filename)
                        if not os.path.exists(save_path):
                            break
                        counter += 1
                        
                    # 3. Save the file
                    with open(save_path, "wb") as f_out:
                        f_out.write(final_image_bytes)
                    print(f"✅ Screenshot copy saved to: {save_path} (Size: {len(final_image_bytes)/1024:.1f} KB)")
                except Exception as e:
                    print(f"⚠️ Failed to save screenshot copy: {e}")

        # --- Generate caption (optional) ---
        ai_caption = None
        now = time.time()

        if SETTINGS.get("captions_enabled", True):
            if now - LAST_CAPTION_TIME >= CAPTION_COOLDOWN_SECONDS:
                LAST_CAPTION_TIME = now
                ai_caption = (
                    caption_image_from_bytes(image_bytes)
                    if input_type == "data"
                    else caption_image(image_input)
                )
                if ai_caption:
                    print(f"✅ Caption generated: '{ai_caption}'")
                else:
                    print("❌ Caption generation returned empty")
            else:
                print("⏳ Caption skipped (cooldown active)")

        # --- Insert image into document ---
        doc = Document(selected_doc_path)

        image_para = None

        if SETTINGS.get("use_codeword_insertion", False):
            image_para = insert_image_at_codeword(doc, final_image_stream)

        if image_para is None:
            doc.add_picture(final_image_stream, width=Inches(6.0))
            image_para = doc.paragraphs[-1]

        # Save document before caption insertion
        try:
            doc.save(selected_doc_path)
        except Exception as e:
            print(f"⚠ Error saving document before caption: {e}")

        # --- Insert caption if available ---
        if SETTINGS.get("captions_enabled", True) and ai_caption:
            prefix = SETTINGS.get("caption_prefix", "Figure")
            if prefix == "Custom":
                prefix = SETTINGS.get("caption_custom_text", "Figure")

            insert_real_caption(selected_doc_path, ai_caption, prefix)
        else:
            if not SETTINGS.get("captions_enabled", True):
                print("📝 Captions disabled by settings.")
            else:
                print("⚠ Screenshot added — no caption (AI empty or rate-limited).")

    except PermissionError:
        app.after(0, handle_permission_error)
    except Exception as e:
        app.after(0, handle_generic_error, repr(e))


# --- Function to process the file after a delay ---
def process_new_image_file(file_path):
    """
    This function runs on the main thread after a 1-second delay,
    giving the screenshot tool time to finish writing the file.
    """
    print(f"Processing file: {file_path}")
    add_image_to_doc(file_path, "path")

# --- Watchdog Event Handler ---
class ImageFileHandler(FileSystemEventHandler):
    def on_created(self, event):
        if not event.is_directory and event.src_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
            print(f"New image detected in folder: {event.src_path}")
            # Add a 1-second delay to prevent file-lock errors
            app.after(1000, process_new_image_file, event.src_path)

# --- Table of Figures Management ---
def update_table_of_figures(doc_path):
    """
    Creates or updates a Table of Figures at the beginning of the document.
    Uses COM (Word Automation) if available for robust handling.
    """
    if not SETTINGS.get("tof_enabled", False):
        print("⏭ TOF disabled in settings. Skipping...")
        return
        
    prefix = SETTINGS.get("caption_prefix", "Figure")
    if prefix == "Custom":
        prefix = SETTINGS.get("caption_custom_text", "Figure")

    if USE_WORD_AUTOMATION:
        print("📑 Updating TOF via Word Automation (COM)...")
        try:
            with COM_LOCK:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(doc_path)
                
                # 1. DELETE EXISTING TOFS (Robust Cleanup)
                # We loop backwards to avoid index shifting issues, theoretically.
                # But treating them as objects works fine.
                count = doc.TablesOfFigures.Count
                if count > 0:
                    print(f"   Found {count} existing tables. Removing...")
                    for i in range(count, 0, -1):
                        doc.TablesOfFigures(i).Delete()

                # 2. DELETE ORPHAN HEADERS
                # Search for "Table of Figures" paragraphs and delete them
                # This cleans up the debris from previous runs
                for p in doc.Paragraphs:
                    # Check text (strip return chars)
                    txt = p.Range.Text.strip() 
                    if txt == "Table of Figures":
                        # Check style if we want to be safe, but text is unique enough usually
                        if "Heading" in p.Style.NameLocal or p.Style.NameLocal == "Heading 1":
                            p.Range.Delete()
                            
                # 3. INSERT NEW TOF
                # Insert after Title (Paragraph 1)
                # If doc is empty, add one.
                if doc.Paragraphs.Count == 0:
                     doc.Content.InsertAfter("Report\r")
                
                # Range strategy: Go to start, move down 1 paragraph
                rng = doc.Paragraphs(1).Range
                rng.Collapse(0) # 0 = wdCollapseEnd
                
                # Insert Header
                rng.InsertParagraphAfter()
                rng_header = doc.Paragraphs(2).Range
                rng_header.Text = "Table of Figures\r"
                rng_header.Style = "Heading 1" # Use internal name "Heading 1"
                
                # Insert TOF Object
                # Move range to after the header
                rng_tof = doc.Paragraphs(3).Range
                rng_tof.Collapse(1) # 1 = wdCollapseStart (Start of para 3)
                
                # Add(Range, Caption, IncludeLabel, UseHeadingStyles, UpperHeadingLevel, LowerHeadingLevel, UseFields, TableID, RightAlignPageNumbers, IncludePageNumbers, AddedStyles, UseHyperlinks, HidePageNumbersInWeb)
                # See MS Docs for Add arguments.
                
                # Simplest Add:
                tof = doc.TablesOfFigures.Add(
                    Range=rng_tof,
                    Caption=prefix,
                    IncludeLabel=True,
                    RightAlignPageNumbers=True,
                    UseHyperlinks=True
                )
                
                # Force update just in case
                tof.Update()
                
                doc.Save()
                doc.Close()
                # word.Quit() # Optional, managed elsewhere usually or let it live?
                word.Quit()
                
            print("✅ TOF successfully rebuilt (COM).")
            return
            
        except Exception as e:
            print(f"⚠️ COM TOF Update failed: {e}. Trying fallback...")
            # If COM fails, fall through to python-docx implementation

    # --- FALLBACK: python-docx (XML Manipulation) ---
    try:
        doc = Document(doc_path)
        
        # Aggressive Cleanup
        entries_to_remove = []
        for i, para in enumerate(doc.paragraphs):
            # Detect Header
            if "Table of Figures" in para.text and para.style.name.startswith("Heading"):
                entries_to_remove.append(para)
                # Also aggressively search for the TOC/Hyperlink styles that follow
                # This is a heuristic: check the Next few paragraphs
                curr_idx = i + 1
                while curr_idx < len(doc.paragraphs):
                    next_p = doc.paragraphs[curr_idx]
                    # If it looks like a TOC entry or empty field
                    if next_p.style.name.startswith("TOC") or next_p.style.name == "Hyperlink" or "Figure" in next_p.text:
                        entries_to_remove.append(next_p)
                        curr_idx += 1
                    else:
                        break
        
        for p in entries_to_remove:
            try:
                p_element = p._element
                if p_element.getparent() is not None:
                    p_element.getparent().remove(p_element)
            except: pass

        doc.save(doc_path)
        doc = Document(doc_path) # Reload

        # Insert Header (Standard)
        if len(doc.paragraphs) > 0:
            if len(doc.paragraphs) > 1:
                p_next = doc.paragraphs[1]
                p_next.insert_paragraph_before("Table of Figures", style="Heading 1")
                tof_field_para = p_next.insert_paragraph_before("") 
            else:
                doc.add_paragraph("Table of Figures", style="Heading 1")
                tof_field_para = doc.add_paragraph("")
        else:
            doc.add_paragraph("Table of Figures", style="Heading 1")
            tof_field_para = doc.add_paragraph("")

        # Insert Field
        instr_text = f' TOC \\h \\z \\c "{prefix}" '
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instr_text
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run = tof_field_para.add_run()
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        doc.save(doc_path)
        print("✅ TOF inserted (Field Codes - Fallback).")
        
    except Exception as e:
        print(f"⚠️ Error updating Table of Figures (Fallback): {e}")

def open_settings_window():
    win = customtkinter.CTkToplevel(app)
    win.title("Settings")
    win.geometry("320x480")
    win.resizable(False, False)
    win.transient(app)
    win.grab_set()

    # -----------------------------
    # Scrollable content area
    # -----------------------------
    scroll = customtkinter.CTkScrollableFrame(
        win,
        width=280,
        height=380
    )
    scroll.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

    # -----------------------------
    # Captions toggle
    # -----------------------------
    var_captions = customtkinter.BooleanVar(
        value=SETTINGS.get("captions_enabled", True)
    )

    chk = customtkinter.CTkCheckBox(
        scroll,
        text="Enable Auto Captions",
        variable=var_captions
    )
    chk.grid(row=0, column=0, padx=10, pady=(10, 8), sticky="w")

    # -----------------------------
    # Caption prefix
    # -----------------------------
    prefixes = ["Figure", "Screenshot", "Step", "Custom"]

    prefix_var = customtkinter.StringVar(
        value=SETTINGS.get("caption_prefix", "Figure")
    )

    pref_label = customtkinter.CTkLabel(
        scroll,
        text="Caption prefix:"
    )
    pref_label.grid(row=1, column=0, padx=10, pady=(6, 2), sticky="w")

    pref_cb = customtkinter.CTkComboBox(
        scroll,
        values=prefixes,
        variable=prefix_var
    )
    pref_cb.grid(row=2, column=0, padx=10, sticky="ew")

    # -----------------------------
    # Custom prefix (only for Custom)
    # -----------------------------
    custom_text_var = customtkinter.StringVar(
        value=SETTINGS.get("caption_custom_text", "")
    )

    custom_entry = customtkinter.CTkEntry(
        scroll,
        textvariable=custom_text_var,
        placeholder_text="Custom prefix"
    )

    def update_custom_prefix_visibility(*_):
        if prefix_var.get() == "Custom":
            custom_entry.configure(state="normal")
            custom_entry.grid(row=3, column=0, padx=10, pady=(6, 10), sticky="ew")
        else:
            custom_entry.grid_remove()
            custom_entry.configure(state="disabled")

    prefix_var.trace_add("write", update_custom_prefix_visibility)
    update_custom_prefix_visibility()

    # -----------------------------
    # Numbering format
    # -----------------------------
    numbering_map = {
        "1, 2, 3": "decimal",
        "a, b, c": "loweralpha",
        "i, ii, iii": "lowerroman",
    }
    reverse_map = {v: k for k, v in numbering_map.items()}

    numbering_var = customtkinter.StringVar(
        value=reverse_map.get(
            SETTINGS.get("numbering_format", "decimal"),
            "1, 2, 3"
        )
    )

    numbering_label = customtkinter.CTkLabel(
        scroll,
        text="Numbering format:"
    )
    numbering_label.grid(row=4, column=0, padx=10, pady=(6, 2), sticky="w")

    numbering_cb = customtkinter.CTkComboBox(
        scroll,
        values=list(numbering_map.keys()),
        variable=numbering_var
    )
    numbering_cb.grid(row=5, column=0, padx=10, sticky="ew")

    # -----------------------------
    # Table of Figures
    # -----------------------------
    var_tof = customtkinter.BooleanVar(
        value=SETTINGS.get("tof_enabled", True)
    )

    tof_chk = customtkinter.CTkCheckBox(
        scroll,
        text="Add / Update Table of Figures",
        variable=var_tof
    )
    tof_chk.grid(row=6, column=0, padx=10, pady=(10, 10), sticky="w")

    # -----------------------------
    # Bottom buttons (fixed)
    # -----------------------------
    btn_frame = customtkinter.CTkFrame(win)
    btn_frame.grid(row=1, column=0, pady=(0, 10))

    def save_and_close():
        SETTINGS["captions_enabled"] = var_captions.get()
        SETTINGS["caption_prefix"] = prefix_var.get()
        SETTINGS["caption_custom_text"] = custom_text_var.get()
        SETTINGS["numbering_format"] = numbering_map[numbering_var.get()]
        SETTINGS["tof_enabled"] = var_tof.get()
        SETTINGS["use_codeword_insertion"] = var_codeword.get()
        SETTINGS["insertion_codeword"] = codeword_var.get()
        SETTINGS["save_copy_enabled"] = var_save_copy.get()
        SETTINGS["save_copy_folder"] = folder_path_var.get()
        SETTINGS["compression_enabled"] = var_compress.get()
        SETTINGS["compression_quality"] = int(comp_slider.get())
             
        save_settings(SETTINGS)
        win.destroy()

    save_btn = customtkinter.CTkButton(
        btn_frame,
        text="Save",
        width=100,
        command=save_and_close
    )
    save_btn.grid(row=0, column=0, padx=6)

    cancel_btn = customtkinter.CTkButton(
        btn_frame,
        text="Cancel",
        width=100,
        command=win.destroy
    )
    cancel_btn.grid(row=0, column=1, padx=6)

    # -----------------------------
    # Codeword insertion
    # -----------------------------
    var_codeword = customtkinter.BooleanVar(
        value=SETTINGS.get("use_codeword_insertion", False)
    )

    codeword_chk = customtkinter.CTkCheckBox(
        scroll,
        text="Insert images at codeword",
        variable=var_codeword
    )
    codeword_chk.grid(row=7, column=0, padx=10, pady=(10, 4), sticky="w")

    codeword_var = customtkinter.StringVar(
        value=SETTINGS.get("insertion_codeword", "[[IMG]]")
    )

    codeword_entry = customtkinter.CTkEntry(
        scroll,
        textvariable=codeword_var,
        placeholder_text="Codeword (e.g. [[IMG]])"
    )

    def update_codeword_visibility(*_):
        if var_codeword.get():
            codeword_entry.configure(state="normal")
            codeword_entry.grid(row=8, column=0, padx=10, pady=(4, 10), sticky="ew")
        else:
            codeword_entry.grid_remove()
            codeword_entry.configure(state="disabled")

    var_codeword.trace_add("write", update_codeword_visibility)
    update_codeword_visibility()

    # -----------------------------
    # Save Copy to Custom Folder
    # -----------------------------
    var_save_copy = customtkinter.BooleanVar(
        value=SETTINGS.get("save_copy_enabled", False)
    )

    save_copy_chk = customtkinter.CTkCheckBox(
        scroll,
        text="Save screenshots to custom folder",
        variable=var_save_copy
    )
    save_copy_chk.grid(row=9, column=0, padx=10, pady=(10, 4), sticky="w")
    
    # Frame for folder selection
    folder_frame = customtkinter.CTkFrame(scroll, fg_color="transparent")
    
    folder_path_var = customtkinter.StringVar(
        value=SETTINGS.get("save_copy_folder", "")
    )
    
    folder_entry = customtkinter.CTkEntry(
        folder_frame,
        textvariable=folder_path_var,
        placeholder_text="No folder selected",
        state="readonly",
        width=200
    )
    folder_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))
    
    def select_save_folder():
        p = filedialog.askdirectory(title="Select Destination Folder")
        if p:
            folder_path_var.set(p)
            
    folder_btn = customtkinter.CTkButton(
        folder_frame,
        text="Browse",
        width=50,
        command=select_save_folder
    )
    folder_btn.pack(side="right")

    def update_save_copy_visibility(*_):
        if var_save_copy.get():
            folder_frame.grid(row=10, column=0, padx=10, pady=(0, 10), sticky="ew")
        else:
            folder_frame.grid_remove()

    var_save_copy.trace_add("write", update_save_copy_visibility)
    update_save_copy_visibility()

    # -----------------------------
    # Image Compression Settings
    # -----------------------------
    var_compress = customtkinter.BooleanVar(
        value=SETTINGS.get("compression_enabled", False)
    )

    compress_chk = customtkinter.CTkCheckBox(
        scroll,
        text="Compress Images (Convert to JPEG)",
        variable=var_compress
    )
    compress_chk.grid(row=11, column=0, padx=10, pady=(10, 4), sticky="w")
    
    # Compression Details Frame
    comp_frame = customtkinter.CTkFrame(scroll, fg_color="transparent")
    
    # Quality Slider
    quality_label_var = customtkinter.StringVar(value=f"Quality: {SETTINGS.get('compression_quality', 85)}")
    
    comp_lbl = customtkinter.CTkLabel(comp_frame, textvariable=quality_label_var)
    comp_lbl.pack(anchor="w")
    
    def update_quality_label(val):
        quality_label_var.set(f"Quality: {int(val)}")

    comp_slider = customtkinter.CTkSlider(
        comp_frame,
        from_=10,
        to=100,
        number_of_steps=90,
        command=update_quality_label
    )
    comp_slider.set(SETTINGS.get("compression_quality", 85))
    comp_slider.pack(fill="x", pady=(0, 10))
    
    def update_compress_visibility(*_):
        if var_compress.get():
            comp_frame.grid(row=12, column=0, padx=20, pady=(0, 10), sticky="ew")
        else:
            comp_frame.grid_remove()
            
    var_compress.trace_add("write", update_compress_visibility)
    update_compress_visibility()


# --- Clipboard Monitoring Loop ---
def check_clipboard_loop():
    global monitoring, LAST_IMAGE_HASH

    if monitoring:
        try:
            img = ImageGrab.grabclipboard()
            if img and isinstance(img, Image.Image): # Ensure it's an image
                # Use raw bytes for stable hashing (PNG compression can vary)
                current_hash = hashlib.md5(img.tobytes()).hexdigest()

                if current_hash != LAST_IMAGE_HASH:
                    LAST_IMAGE_HASH = current_hash
                    print(f"New image detected! (Hash: {current_hash[:8]}...)")
                    
                    # Convert to PNG bytes for insertion
                    with io.BytesIO() as output:
                        img.save(output, format="PNG")
                        png_data = output.getvalue()
                        
                    add_image_to_doc(png_data, input_type="data")
            
            # If clipboard is empty or not an image (e.g. text/files), we ignore or handle?
            # Current logic just ignores.

        except Exception as e:
            print(f"Error checking clipboard: {e}")

        app.after(1200, check_clipboard_loop)

# --- GUI Functions ---
def browse_for_file():
    global selected_doc_path
    path = filedialog.askopenfilename(
        title="Select your Lab Report (.docx)",
        filetypes=[("Word documents", "*.docx")]
    )
    if path:
        selected_doc_path = path
        filename = os.path.basename(path)
        file_label.configure(text=f"File: {filename}", text_color=COLOR_TEXT_IDLE)
    else:
        selected_doc_path = None
        file_label.configure(text="No file selected.", text_color=COLOR_TEXT_ERROR)

def browse_for_folder():
    global selected_folder_path
    path = filedialog.askdirectory(title="Select Folder to Monitor")
    if path:
        selected_folder_path = path
        folder_label.configure(text=f"Folder: {os.path.basename(path)}", text_color=COLOR_TEXT_IDLE)
    else:
        selected_folder_path = None
        folder_label.configure(text="No folder selected.", text_color=COLOR_TEXT_ERROR)

def select_monitoring_mode(mode):
    global monitoring_mode, selected_folder_path
    monitoring_mode = mode
    print(f"Mode switched to: {mode}")
    if mode == "Clipboard":
        browse_folder_button.configure(state="disabled")
        folder_label.configure(text="")
    else: # Folder mode
        browse_folder_button.configure(state="normal")
        if selected_folder_path:
            folder_label.configure(text=f"Folder: {os.path.basename(selected_folder_path)}", text_color=COLOR_TEXT_IDLE)
        else:
            folder_label.configure(text="No folder selected.", text_color=COLOR_TEXT_ERROR)

def start_monitoring():
    global monitoring, monitoring_mode, folder_observer, selected_folder_path
    
    if not selected_doc_path:
        messagebox.showerror("Error", "Please select a .docx file first!")
        return
    
    monitoring = True
    
    if monitoring_mode == "Folder":
        if not selected_folder_path:
            messagebox.showerror("Error", "Please select a folder to monitor first!")
            monitoring = False 
            return
        
        print(f"Starting folder monitoring for: {selected_folder_path}")
        event_handler = ImageFileHandler()
        folder_observer = Observer()
        folder_observer.schedule(event_handler, selected_folder_path, recursive=False)
        folder_observer.start() 
        
    else: # Clipboard mode
        print("Starting clipboard monitoring...")
        
        # 1. Clear clipboard to strictly enforce "only images taken AFTER start"
        try:
            app.clipboard_clear()
            print("🧹 Clipboard cleared. Ready for NEW screenshots.")
        except Exception as e:
            print(f"⚠️ Could not clear clipboard: {e}")

        # 2. Reset Hash
        LAST_IMAGE_HASH = None
        
        check_clipboard_loop() 
    
    start_button.configure(state="disabled")
    stop_button.configure(state="normal")
    browse_button.configure(state="disabled")
    browse_folder_button.configure(state="disabled") 
    mode_switch.configure(state="disabled") 
    file_label.configure(text_color=COLOR_TEXT_ACTIVE) 

def stop_monitoring():
    global monitoring, monitoring_mode, folder_observer
    
    print("Monitoring stopped.")
    monitoring = False 
    
    if monitoring_mode == "Folder" and folder_observer:
        print("Stopping folder observer...")
        folder_observer.stop()
        folder_observer.join() 
        folder_observer = None
        print("Folder observer stopped.")

    # --- NEW: Update Table of Figures before stopping ---
    if SETTINGS.get("tof_enabled", False) and selected_doc_path:
        print("📑 Updating Table of Figures...")
        update_table_of_figures(selected_doc_path)
    
    start_button.configure(state="normal")
    stop_button.configure(state="disabled")
    browse_button.configure(state="normal")
    mode_switch.configure(state="normal") 
    file_label.configure(text_color=COLOR_TEXT_IDLE)
    
    if monitoring_mode == "Folder":
        browse_folder_button.configure(state="normal")

def on_closing():
    global monitoring, folder_observer
    print("Application closing...")
    
    # --- NEW: Update Table of Figures before closing ---
    if SETTINGS.get("tof_enabled", False) and selected_doc_path:
        print("📑 Updating Table of Figures before exit...")
        update_table_of_figures(selected_doc_path)
    
    save_settings(SETTINGS)
    monitoring = False
    if folder_observer:
        folder_observer.stop()
        folder_observer.join()
    app.destroy()


# --- UI SETUP ---
customtkinter.set_appearance_mode("light")

# --- Load Icons ---
try:
    # --- MODIFIED: Use the resource_path function and your new filename "doc.png" ---
    icon_browse_path = resource_path("icons/doc.png")
    icon_browse = customtkinter.CTkImage(
        light_image=Image.open(icon_browse_path), size=(48, 48)) 
except FileNotFoundError:
    print("doc.png not found. Running without icon.")
    icon_browse = None

# --- Create Window ---
app = customtkinter.CTk()
app.title("DocuDrop")
app.geometry("400x580") 
app.configure(fg_color=COLOR_WINDOW_BG)
app.protocol("WM_DELETE_WINDOW", on_closing)
app.minsize(400, 580) 
app.maxsize(400, 580) # Added maxsize as per your code

# load gear icon (add icons/gear.png to resources)
try:
    gear_path = resource_path("icons/gear.png")
    gear_pil = Image.open(gear_path).convert("RGBA")  # force correct mode
    gear_img = customtkinter.CTkImage(
        light_image=gear_pil,
        size=(26, 26)   # slightly larger for clarity
    )
except Exception as e:
    print("⚠ Settings icon not loaded:", e)
    gear_img = None

top_bar = customtkinter.CTkFrame(app, fg_color="transparent")
top_bar.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(6, 0))
top_bar.grid_columnconfigure(0, weight=1)

settings_button = customtkinter.CTkButton(
    master=top_bar,
    text="",                     # icon only
    image=gear_img,
    width=34,
    height=34,
    fg_color="transparent",
    hover_color=COLOR_BUTTON,
    corner_radius=8,
    command=open_settings_window
)
settings_button.grid(row=0, column=1, padx=10, sticky="e")

try:
    # --- MODIFIED: Use the resource_path function ---
    icon_path = resource_path("icons/app_icon.ico")
    app.iconbitmap(icon_path)
    print("App icon loaded successfully.")
except Exception as e:
    print(f"Could not load app icon: {e}. Must be a .ico file.")

# --- Configure the main grid layout ---
app.grid_rowconfigure(0, weight=0) 
app.grid_rowconfigure(1, weight=0) 
app.grid_rowconfigure(2, weight=0) # Icon Label
app.grid_rowconfigure(3, weight=0) # Browse Button
app.grid_rowconfigure(4, weight=0) # File Label
app.grid_rowconfigure(5, weight=1) # Mode Switch
app.grid_rowconfigure(6, weight=0) # Browse Folder
app.grid_rowconfigure(7, weight=0) # Folder Label
app.grid_rowconfigure(8, weight=1) # Start/Stop
app.grid_columnconfigure(0, weight=1) 
app.grid_columnconfigure(1, weight=1) 

# --- Define Fonts ---
title_font = customtkinter.CTkFont(family="Forte", size=45, weight="bold")
welcome_font = customtkinter.CTkFont(family="Forte", size=20)
button_font = customtkinter.CTkFont(family="Forte", size=25, weight="bold") 
label_font = customtkinter.CTkFont(family="Forte", size=25)
small_label_font = customtkinter.CTkFont(family="Forte", size=16)
switch_font = customtkinter.CTkFont(family="Forte", size=22, weight="bold")


# --- Title and Welcome Labels ---
title_label = customtkinter.CTkLabel(master=app, text="DocuDrop", text_color=COLOR_TEXT, font=title_font)
title_label.grid(row=0, column=0, columnspan=2, pady=(20, 0), padx=10) 

welcome_label = customtkinter.CTkLabel(master=app, text="Welcome! Select your file to begin.", text_color=COLOR_TEXT_IDLE, font=welcome_font)
welcome_label.grid(row=1, column=0, columnspan=2, pady=(0, 10), padx=10)

# --- Standalone Icon Label ---
icon_label = customtkinter.CTkLabel(master=app, text="", image=icon_browse)
icon_label.grid(row=2, column=0, columnspan=2, pady=(10, 0), padx=10)

# --- Docx Widgets ---
browse_button = customtkinter.CTkButton(
    master=app,
    text="Select .docx File",
    command=browse_for_file,
    fg_color=COLOR_BUTTON,
    hover_color=COLOR_BUTTON_HOVER,
    text_color=COLOR_TEXT,
    font=button_font,
    text_color_disabled=COLOR_TEXT_DISABLED,
    corner_radius=10,
    height=50
)
browse_button.grid(row=3, column=0, columnspan=2, pady=(5, 10), padx=40) 

file_label = customtkinter.CTkLabel(master=app, text="No file selected", text_color=COLOR_TEXT_ERROR, font=label_font)
file_label.grid(row=4, column=0, columnspan=2, pady=(0, 10), padx=10)

# --- Monitoring Mode UI ---
mode_switch = customtkinter.CTkSegmentedButton(
    master=app,
    values=["Clipboard", "Folder"],
    command=select_monitoring_mode,
    font=switch_font,
    text_color=COLOR_TEXT, 
    selected_color=COLOR_BUTTON,
    selected_hover_color=COLOR_BUTTON_HOVER,
    unselected_color=COLOR_WINDOW_BG,
    unselected_hover_color=COLOR_BUTTON,
    border_width=2,
    fg_color=COLOR_WINDOW_BG,
)
mode_switch.grid(row=5, column=0, columnspan=2, pady=10, padx=40, sticky="ew")
mode_switch.set("Clipboard") # Set default

browse_folder_button = customtkinter.CTkButton(
    master=app,
    text="Select Folder",
    command=browse_for_folder,
    fg_color=COLOR_BUTTON,
    hover_color=COLOR_BUTTON_HOVER,
    text_color=COLOR_TEXT,
    font=button_font,
    text_color_disabled=COLOR_TEXT_DISABLED,
    corner_radius=10,
    height=40,
    state="disabled" # Starts disabled
)
browse_folder_button.grid(row=6, column=0, columnspan=2, pady=5, padx=40)

folder_label = customtkinter.CTkLabel(master=app, text="", text_color=COLOR_TEXT_IDLE, font=small_label_font)
folder_label.grid(row=7, column=0, columnspan=2, pady=(0, 10), padx=10)


# --- Start/Stop buttons ---
start_button = customtkinter.CTkButton(
    master=app,
    text="Start",
    command=start_monitoring,
    fg_color=COLOR_BUTTON,
    hover_color=COLOR_BUTTON_HOVER,
    text_color=COLOR_TEXT,
    font=button_font,
    text_color_disabled=COLOR_TEXT_DISABLED,
    corner_radius=10,
    height=50
)


stop_button = customtkinter.CTkButton(
    master=app,
    text="Stop",
    command=stop_monitoring,
    state="disabled",
    fg_color=COLOR_BUTTON,
    hover_color=COLOR_BUTTON_HOVER,
    text_color=COLOR_TEXT, 
    text_color_disabled=COLOR_TEXT_DISABLED,
    font=button_font,
    corner_radius=10,
    height=50
)
start_button.grid(row=8, column=0, pady=20, padx=20)
stop_button.grid(row=8, column=1, pady=20, padx=20)

# --- Start the Application ---
print("DocuDrop v11 (EXE-Ready) started. Waiting for user interaction.")
app.mainloop()
print("Application closed.")
